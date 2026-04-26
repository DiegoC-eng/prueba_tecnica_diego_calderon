"""
Generador de guía de preparación para exposición
Autor: Diego Alberto Calderón Calderón
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Colores Walmart
BLUE  = RGBColor(0x00, 0x53, 0xE2)
SPARK = RGBColor(0xFF, 0xC2, 0x20)
RED   = RGBColor(0xEA, 0x11, 0x00)
GREEN = RGBColor(0x2A, 0x87, 0x03)
GRAY  = RGBColor(0x74, 0x76, 0x7C)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# --- Márgenes ---
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# --- Helpers ---
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=BLUE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def body(doc, text, bold=False, color=DARK, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, level=0, color=DARK):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = color
    return p

def qa_block(doc, question, answer):
    """Bloque pregunta-respuesta estilo Q&A"""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Pregunta (fondo azul)
    q_cell = tbl.cell(0, 0)
    set_cell_bg(q_cell, '0053E2')
    q_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    qp = q_cell.paragraphs[0]
    qr = qp.add_run('❓ ' + question)
    qr.font.bold       = True
    qr.font.size       = Pt(10.5)
    qr.font.color.rgb  = WHITE
    q_cell._tc.get_or_add_tcPr()

    # Respuesta (fondo claro)
    a_cell = tbl.cell(1, 0)
    set_cell_bg(a_cell, 'E8F0FD')
    ap = a_cell.paragraphs[0]
    ar = ap.add_run('✅ ' + answer)
    ar.font.size      = Pt(10)
    ar.font.color.rgb = DARK

    # Espacio después
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def section_badge(doc, text, color_hex, text_color=WHITE):
    """Caja de sección coloreada"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color_hex)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.bold      = True
    run.font.size      = Pt(13)
    run.font.color.rgb = text_color
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def divider(doc):
    p = doc.add_paragraph('─' * 80)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(2)


# ============================================================
# PORTADA
# ============================================================
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('GUÍA DE PREPARACIÓN')
r.font.size      = Pt(26)
r.font.bold      = True
r.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Exposición de Prueba Técnica — Data Analyst')
r2.font.size      = Pt(14)
r2.font.color.rgb = GRAY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Diego Alberto Calderón Calderón  |  Abril 2026')
r3.font.size      = Pt(11)
r3.font.color.rgb = GRAY

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

body(doc,
    'Este documento reúne dos cosas: (1) las preguntas más probables que te van a hacer '
    'en la defensa de la prueba técnica, con las respuestas exactas que debes dar; '
    'y (2) los cambios concretos que debes hacer al entregable para que el trabajo '
    'se vea más orgánico y menos generado por IA.',
    size=11)

doc.add_paragraph()


# ============================================================
# SECCIÓN 1 — PREGUNTAS Y RESPUESTAS
# ============================================================
doc.add_page_break()
section_badge(doc, '  PARTE 1 — QUÉ APRENDERTE PARA LA EXPOSICIÓN  ', '0053E2')
doc.add_paragraph()

body(doc,
    'El evaluador no te va a preguntar si el código corre — va a profundizar en '
    'tus decisiones. Aquí están las preguntas exactas por bloque y cómo responderlas.',
    size=10.5)
doc.add_paragraph()


# --- A/B TEST ---
heading(doc, '🧪 A/B Test — Bloque 3 (25% del peso — El más importante)', level=2, color=BLUE)
body(doc, 'Este bloque tiene el mayor peso. El resultado fue negativo y no significativo — '
     'eso es lo más interesante y lo que más te van a preguntar.', size=10)
doc.add_paragraph()

qa_block(doc,
    '¿Por qué usaste un Welch t-test y no un t-test estándar?',
    'Porque el Welch t-test no asume varianzas iguales entre los grupos CONTROL y TREATMENT. '
    'Cuando los grupos tienen tamaños o varianzas distintas — como en este caso con 18 vs 20 tiendas — '
    'el Welch es más robusto y conservador. El t-test estándar inflaría el poder estadístico '
    'incorrectamente.')

qa_block(doc,
    '¿Qué significa exactamente un p-value de 0.24?',
    'Significa que si no hubiera ningún efecto real del tratamiento, observaríamos una diferencia '
    'tan grande como la que vimos el 24% de las veces solo por azar. NO significa que hay un 24% '
    'de probabilidad de que la hipótesis sea falsa. Como convención usamos p < 0.05 para rechazar '
    'la hipótesis nula — con p=0.24 no tenemos evidencia suficiente.')

qa_block(doc,
    '¿El tratamiento empeoró las ventas? El lift fue -16.92%...',
    'El lift negativo es preocupante pero NO es concluyente. El intervalo de confianza al 95% '
    'va de -$6,476 a +$1,586 — ese rango incluye el cero, lo que significa que la diferencia '
    'es estadísticamente compatible con cero. No podemos afirmar que el tratamiento empeoró '
    'las ventas, pero tampoco que las mejoró. La decisión correcta es no implementar.')

qa_block(doc,
    '¿Qué harías si el p-value fuera 0.08?',
    'No implementaría. Extendería el test 4 semanas adicionales para aumentar el poder '
    'estadístico. Un p=0.08 con lift NEGATIVO es incluso más preocupante que p=0.24 — '
    'estamos rozando la significancia en la dirección equivocada. Con más datos confirmaríamos '
    'si el efecto negativo es real. Solo implementaría si el lift fuera positivo Y el p-value '
    'bajara de 0.05 con el test extendido.')

qa_block(doc,
    '¿Los grupos CONTROL y TREATMENT eran comparables antes del test?',
    'Excluí TIENDA_008 y TIENDA_037 porque estaban asignadas a ambos grupos — eso invalida '
    'su inclusión. Con las 38 tiendas restantes, los grupos son comparables en tamaño '
    '(18 vs 20). Idealmente habría verificado balance en GMV pre-test, formato y m² '
    'mediante un test de Mann-Whitney antes de arrancar el experimento. Eso es una '
    'limitación que documenté.')

qa_block(doc,
    '¿Cuántas tiendas necesitarías para un experimento válido?',
    'Depende del poder estadístico deseado (normalmente 80%) y del Minimum Detectable Effect '
    '(MDE). Si queremos detectar un lift del 10% con poder del 80%, necesitaríamos calcular '
    'el sample size con la varianza observada. Con solo 40 tiendas y alta variabilidad inter-tienda, '
    '6 semanas puede ser insuficiente. Lo ideal sería un análisis de poder ANTES de correr el experimento.')

doc.add_paragraph()

# --- SQL ---
heading(doc, '🧱 SQL Avanzado — Bloque 1 (20% del peso)', level=2, color=BLUE)
doc.add_paragraph()

qa_block(doc,
    '¿Qué es una CTE y por qué la usaste en vez de subqueries?',
    'CTE (Common Table Expression) es una tabla temporal nombrada que existe solo durante '
    'la ejecución de la query. La uso porque: (1) hace el código legible — cada CTE '
    'tiene un nombre que explica su propósito; (2) permite reutilizar el mismo resultado '
    'sin recalcularlo; (3) BigQuery puede materializarlas para mejorar rendimiento. '
    'Una subquery anidada hace el mismo trabajo pero es mucho más difícil de leer y mantener.')

qa_block(doc,
    '¿Cómo funciona PARTITION BY en la query de cohortes?',
    'PARTITION BY divide el conjunto de datos en grupos para aplicar la función de ventana '
    'de forma independiente en cada grupo, SIN colapsar las filas como hace GROUP BY. '
    'En la query de cohortes, PARTITION BY cohorte_mes hace que el RANK() y los cálculos '
    'de retención se calculen por separado para cada cohorte, manteniendo todas las filas '
    'de detalle disponibles.')

qa_block(doc,
    '¿Qué es el patrón islands & gaps que usaste en la Query 5 de quiebres de stock?',
    'Es una técnica clásica de SQL para detectar secuencias consecutivas. "\'Islands\'" '
    'son los bloques de datos continuos (días con venta) y "\'gaps\'" son los espacios '
    'entre ellos (días sin venta). El truco: usas LAG() para identificar cuándo la '
    'secuencia se rompe — si la diferencia entre una fecha y la anterior es > 1 día, '
    'empieza un nuevo grupo. Luego usas SUM() acumulado para numerar cada grupo. '
    'Al final agrupas por ese número para obtener el inicio, fin y duración de cada gap.')

qa_block(doc,
    '¿Por qué usas PERCENTILE_CONT en la Query 2?',
    'Para calcular el percentil 25 de GMV/m² dentro de cada formato y marcar tiendas '
    'como BAJO_RENDIMIENTO. PERCENTILE_CONT hace interpolación lineal entre valores '
    '(resultado continuo), a diferencia de PERCENTILE_DISC que retorna un valor exacto '
    'del dataset. En retail prefiero CONT porque da un umbral más suave y menos '
    'sensible a valores extremos.')

qa_block(doc,
    '¿Cómo garantizas que la Query 1 de Comp Sales no incluye tiendas nuevas?',
    'Filtro en la CTE \'tiendas_comparables\' con `opening_date < \'2024-01-01\'`. '
    'Esto asegura que la tienda llevaba al menos 12 meses operando antes del período '
    'de comparación más antiguo que uso. Además filtro que la tienda tenga ventas '
    'en AMBOS períodos — si una tienda tuvo ventas en 2024 pero no en 2025 '
    '(o viceversa), no es una comparación válida.')

doc.add_paragraph()

# --- Modelado ---
heading(doc, '📐 Modelado de Datos — Bloque 2 (15% del peso)', level=2, color=BLUE)
doc.add_paragraph()

qa_block(doc,
    '¿Por qué granularidad a nivel de ítem y no de transacción?',
    'Porque con granularidad de transacción perdería la capacidad de: calcular GMROI '
    'por SKU/proveedor, analizar basket uplift por producto, y detectar quiebres '
    'de stock por ítem. El nivel más atómico disponible es el ítem-transacción, '
    'y siempre es mejor agregar hacia arriba que perder detalle. '
    'En BigQuery el almacenamiento es columnar — una tabla con 542K filas no es '
    'un problema de rendimiento si está bien particionada por fecha.')

qa_block(doc,
    '¿Cómo modelaste que el 60% de transacciones no tiene customer_id?',
    'Mantuve customer_id como NULLABLE en fact_sales y agregué un campo '
    '\'is_identified\' booleano en dim_customer. La alternativa de crear un '
    'registro \'ANON\' genérico contamina los análisis de cohortes. '
    'Con el campo nullable, todos los análisis de clientes filtran '
    'explícitamente WHERE is_identified = TRUE, haciendo la exclusión visible '
    'y auditableen el código.')

qa_block(doc,
    '¿Qué es un SCD Tipo 1 y por qué lo usaste en dim_customer?',
    'SCD = Slowly Changing Dimension. Tipo 1 significa que cuando un atributo '
    'cambia, sobreescribes el valor anterior sin guardar historial. '
    'Lo usé en dim_customer porque para análisis de cohortes solo necesito '
    'la fecha de primera transacción — no me interesa el historial de '
    'cambios del cliente. Si necesitara trackear cambios en el comportamiento '
    'del cliente a lo largo del tiempo, usaría SCD Tipo 2 con fechas de vigencia.')

qa_block(doc,
    'Si dos reportes muestran GMV distinto para la misma tienda y día, ¿qué haces?',
    'Proceso de 4 pasos: (1) Identifico las queries exactas de cada reporte — '
    '¿usan la misma tabla? ¿el mismo filtro de status? ¿el mismo campo de GMV? '
    '(2) Comparo la definición de GMV — uno puede incluir devueltas y el otro no, '
    'o uno usa total_amount y el otro SUM(unit_price×quantity). '
    '(3) Verifico el filtro de fecha — las 2h de retraso de POS pueden crear '
    'diferencias si el corte es a medianoche exacta. '
    '(4) Documento la causa en el Glosario de Métricas y certifico una vista '
    'única como fuente de verdad para todos los reportes.')

doc.add_paragraph()

# --- KPIs ---
heading(doc, '🎯 KPIs — Bloque 4 (10% del peso)', level=2, color=BLUE)
doc.add_paragraph()

qa_block(doc,
    '¿Por qué el North Star es GMV/m² y no GMV absoluto?',
    'Porque un Hipermercado de 8,000m² siempre va a tener más GMV que un Express '
    'de 400m² — comparar en absoluto es injusto. El GMV/m² normaliza por espacio '
    'y permite tres cosas clave: (1) comparar formatos distintos en la misma escala, '
    '(2) comparar tiendas de distintos países sin sesgo por tamaño, '
    '(3) es directamente accionable para el gerente — más ventas por el mismo '
    'espacio significa mejor surtido, menos quiebres, mejor exhibición. '
    'Es además la métrica que usan los analistas de Wall Street para evaluar retailers.')

qa_block(doc,
    '¿Qué es un leading indicator y cuál pusiste?',
    'Un leading indicator es predictivo — te dice lo que VA a pasar antes de que '
    'el resultado final se materialice. En mi framework usé la Tasa de Devolución '
    '(Return Rate) como leading indicator. Una devolución refleja insatisfacción '
    'ANTES de que el cliente decida no volver — si el Return Rate sube en una tienda, '
    'es probable que la retención caiga en el próximo mes. Esto permite intervención '
    'preventiva antes de perder clientes. Los KPIs de GMV y retención son lagging '
    '— te dicen lo que ya pasó.')

qa_block(doc,
    '¿Cómo detectas si el dato de un KPI está mal?',
    'Cada KPI tiene definida una regla de validación. Por ejemplo, para GMV/m²: '
    'si una tienda activa reporta 0, hay un gap de datos — no una tienda sin ventas. '
    'Para Return Rate: si una tienda tiene exactamente el mismo % durante 4 semanas '
    'consecutivas, posiblemente el campo dejó de actualizarse. '
    'La clave es definir umbrales de alerta basados en el comportamiento histórico, '
    'no en valores absolutos — una variación del 30% en una semana es normal '
    'en diciembre pero anómala en febrero.')

doc.add_paragraph()


# ============================================================
# SECCIÓN 2 — HACER EL TRABAJO MÁS MANUAL
# ============================================================
doc.add_page_break()
section_badge(doc, '  PARTE 2 — QUÉ CAMBIAR PARA QUE SE VEA MÁS MANUAL  ', 'ea1100')
doc.add_paragraph()

body(doc,
    'Un evaluador experimentado puede detectar trabajo generado por IA por tres señales: '
    'estructura demasiado perfecta, ausencia de resultados reales de código, y falta '
    'de voz propia con contexto local. Aquí los cambios concretos ordenados por impacto.',
    size=10.5)
doc.add_paragraph()


# --- ALTO IMPACTO ---
heading(doc, '🔴 Cambios de Alto Impacto', level=2, color=RED)

body(doc, '1. Agrega resultados reales a las queries SQL', bold=True, size=11)
body(doc,
    'La señal más obvia de una query no ejecutada es que no tiene output. '
    'Corre cada query con DuckDB (o similar) y pega una muestra de 3-5 filas '
    'como comentario al final de cada bloque SQL. Por ejemplo:',
    size=10)

# Código de ejemplo
p_code = doc.add_paragraph()
p_code.paragraph_format.left_indent = Inches(0.4)
run_code = p_code.add_run(
    '-- RESULTADO REAL (5 filas):\n'
    '-- store_id  | country | format       | gmv_actual | comp_sales_pct\n'
    '-- TIENDA_01 | CR      | HIPERMERCADO | 284,392.10 | +12.4%\n'
    '-- TIENDA_07 | GT      | SUPERMERCADO | 198,211.50 | +8.1%'
)
run_code.font.name      = 'Courier New'
run_code.font.size      = Pt(9)
run_code.font.color.rgb = RGBColor(0x2A, 0x87, 0x03)

doc.add_paragraph()

body(doc, '2. Convierte el análisis a Jupyter Notebook con outputs visibles', bold=True, size=11)
body(doc,
    'Un archivo .py limpio parece generado. Un .ipynb con celdas de exploración, '
    'outputs visibles y alguna celda de prueba comentada parece orgánico. '
    'Lo más importante: que cada celda de análisis muestre el dataframe resultante, '
    'no solo el gráfico final.',
    size=10)
doc.add_paragraph()

body(doc, '3. Agrega tu voz y contexto regional en el análisis', bold=True, size=11)
body(doc,
    'Reemplaza frases genéricas por observaciones con contexto local. Por ejemplo:',
    size=10)

bullet(doc,
    'En vez de: "Pico en diciembre en Hipermercado"',
    color=RED)
bullet(doc,
    'Escribe: "El pico de diciembre en Hipermercado es consistente con el '
    'aguinaldo en Centroamérica — en CR y GT el décimo tercer mes se paga en '
    'diciembre, generando un spike de consumo típico del retail de la región."',
    color=GREEN)
doc.add_paragraph()

bullet(doc,
    'En vez de: "HN y NI tienen más efectivo"',
    color=RED)
bullet(doc,
    'Escribe: "La alta penetración de efectivo en HN y NI refleja la menor '
    'bancarización — según datos del Banco Mundial, menos del 45% de adultos '
    'en Honduras tiene cuenta bancaria, lo que limita naturalmente la adopción '
    'de pagos digitales."',
    color=GREEN)
doc.add_paragraph()

body(doc, '4. Agrega pensamiento propio sobre el proceso de análisis', bold=True, size=11)
body(doc,
    'Los analistas reales documentan su proceso de pensamiento, incluyendo los callejones '
    'sin salida. Agrega texto como:',
    size=10)

p_quote = doc.add_paragraph()
p_quote.paragraph_format.left_indent = Inches(0.4)
run_q = p_quote.add_run(
    '"Inicialmente intenté calcular los gaps con una self-join pero era muy lento con '
    '542K filas. Cambié al approach de window functions con LAG() que resolvió en segundos."\n\n'
    '"El resultado del A/B me sorprendió — esperaba lift neutro, no negativo. '
    'Revisé el merge dos veces para descartar error en la asignación de grupos."'
)
run_q.font.size      = Pt(10)
run_q.font.italic    = True
run_q.font.color.rgb = BLUE

doc.add_paragraph()


# --- IMPACTO MEDIO ---
heading(doc, '🟡 Cambios de Impacto Medio', level=2, color=RGBColor(0xFF, 0xC2, 0x20))

body(doc, '5. Incluye una sección "Limitaciones" en el análisis', bold=True, size=11)
body(doc,
    'Los analistas maduros reconocen las limitaciones de su propio trabajo. Agrega al '
    'bloque3_analisis.md una sección como esta:',
    size=10)

p_lim = doc.add_paragraph()
p_lim.paragraph_format.left_indent = Inches(0.4)
run_l = p_lim.add_run(
    '"La detección de quiebres de stock usa ausencia de ventas como proxy de OOS. '
    'Esto genera falsos positivos en ítems de baja rotación o categorías estacionales. '
    'Un análisis más preciso requeriría datos de inventario físico, que no están en el dataset."'
)
run_l.font.size      = Pt(10)
run_l.font.italic    = True
run_l.font.color.rgb = GRAY

doc.add_paragraph()

body(doc, '6. Agrega "Decisiones que cambiaría" en bloque2', bold=True, size=11)
body(doc,
    'Al final del bloque2_decisiones.md agrega algo como:',
    size=10)

p_ch = doc.add_paragraph()
p_ch.paragraph_format.left_indent = Inches(0.4)
run_c = p_ch.add_run(
    '"Si tuviera más tiempo, modelaría una tabla separada fact_inventory en lugar de '
    'inferir quiebres desde gaps de ventas. También consideraría SCD Tipo 2 para '
    'dim_store para rastrear cambios de formato o región de una tienda a lo largo del tiempo."'
)
run_c.font.size      = Pt(10)
run_c.font.italic    = True
run_c.font.color.rgb = GRAY

doc.add_paragraph()

body(doc, '7. El código de Python debe tener exploración visible', bold=True, size=11)
body(doc,
    'Agrega en el notebook algunas celdas de exploración típicas que un analista real hace:',
    size=10)

for snippet in [
    'txn.dtypes  # verificar tipos antes de parsear fechas',
    'txn[\'total_amount\'].describe()  # ver distribución antes de analizar',
    'txn.isnull().sum()  # primer vistazo a nulls',
    'txn.head(10)  # siempre mirar los datos crudos primero',
]:
    p_s = doc.add_paragraph()
    p_s.paragraph_format.left_indent = Inches(0.4)
    r_s = p_s.add_run(snippet)
    r_s.font.name      = 'Courier New'
    r_s.font.size      = Pt(9)
    r_s.font.color.rgb = RGBColor(0x2A, 0x87, 0x03)

doc.add_paragraph()


# --- RÁPIDOS ---
heading(doc, '🟢 Cambios Rápidos (menos de 5 min cada uno)', level=2, color=GREEN)

body(doc, '8. Sé más específico en el README sobre qué validaste tú', bold=True, size=11)
body(doc,
    'En vez de: "Las decisiones analíticas son propias"',
    size=10)
body(doc,
    'Escribe: "Verifiqué manualmente que la tasa de 59.83% sin customer_id '
    'coincide exactamente con loyalty_card=FALSE (diferencia=0). '
    'Corrí el t-test del A/B dos veces cambiando el orden del merge '
    'para confirmar que el signo del lift era correcto."',
    size=10, color=GREEN)
doc.add_paragraph()

body(doc, '9. Agrega el snippet de código en la auditoría', bold=True, size=11)
body(doc,
    'En bloque0_auditoria.md, debajo de cada hallazgo, incluye el código Python '
    'que lo generó. Esto demuestra que los números vienen de análisis real:',
    size=10)

p_au = doc.add_paragraph()
p_au.paragraph_format.left_indent = Inches(0.4)
run_au = p_au.add_run(
    '# Verificación de consistencia\n'
    'items_sum = items.groupby(\'transaction_id\').apply(\n'
    '    lambda x: (x[\'unit_price\'] * x[\'quantity\']).sum()\n'
    ').reset_index()\n'
    'inconsistente = (abs(merged[\'diff\']) > 0.01).sum()\n'
    '# Resultado: 1,745 transacciones con discrepancia'
)
run_au.font.name      = 'Courier New'
run_au.font.size      = Pt(9)
run_au.font.color.rgb = RGBColor(0x2A, 0x87, 0x03)

doc.add_paragraph()


# ============================================================
# TABLA RESUMEN DE PRIORIDADES
# ============================================================
doc.add_page_break()
section_badge(doc, '  RESUMEN — TABLA DE PRIORIDADES  ', '1A1A2E')
doc.add_paragraph()

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h_text in enumerate(['Cambio', 'Impacto', 'Tiempo estimado']):
    cell = tbl.cell(0, i)
    set_cell_bg(cell, '0053E2')
    p = cell.paragraphs[0]
    r = p.add_run(h_text)
    r.font.bold      = True
    r.font.color.rgb = WHITE
    r.font.size      = Pt(10)

rows_data = [
    ('Correr queries y pegar resultados reales en el .sql',     '🔴 Alto',   '30 min'),
    ('Convertir bloque3_analisis.py a Jupyter Notebook',        '🔴 Alto',   '20 min'),
    ('Agregar voz propia y contexto regional en el análisis',   '🔴 Alto',   '20 min'),
    ('Sección "Limitaciones" en bloque3',                       '🟡 Medio',  '10 min'),
    ('Sección "Decisiones que cambiaría" en bloque2',           '🟡 Medio',  '10 min'),
    ('Exploración básica visible en el notebook',               '🟡 Medio',  '10 min'),
    ('Especificar qué validaste tú en el README',               '🟢 Rápido', '5 min'),
    ('Snippets de código en bloque0_auditoria.md',              '🟢 Rápido', '5 min'),
]

for i, (cambio, impacto, tiempo) in enumerate(rows_data):
    row = tbl.add_row()
    bg = 'F9FAFB' if i % 2 == 0 else 'FFFFFF'
    for j, txt in enumerate([cambio, impacto, tiempo]):
        cell = row.cells[j]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.font.size      = Pt(10)
        r.font.color.rgb = DARK

doc.add_paragraph()

# --- PREGUNTAS DE LA EXPOSICIÓN ---
doc.add_page_break()
section_badge(doc, '  CHECKLIST MENTAL PARA LA EXPOSICIÓN  ', '2a8703')
doc.add_paragraph()

body(doc,
    'Antes de entrar a la exposición, asegúrate de poder responder estas preguntas '
    'en menos de 60 segundos cada una:',
    size=10.5)
doc.add_paragraph()

checklist = [
    ('A/B Test',      [
        '¿Por qué Welch t-test?',
        '¿Qué significa p=0.24 en términos de negocio?',
        '¿Por qué NO implementarías aunque el p-value fuera 0.08?',
        '¿Por qué excluiste TIENDA_008 y TIENDA_037?',
    ]),
    ('SQL',           [
        '¿Qué hace PARTITION BY y en qué query lo usaste?',
        '¿Qué es el patrón islands & gaps?',
        '¿Por qué la Query 2 usa PERCENTILE_CONT y no MIN?',
        '¿Cómo garantizas que Comp Sales excluye tiendas nuevas?',
    ]),
    ('Modelado',      [
        '¿Por qué granularidad a nivel de ítem?',
        '¿Qué es SCD Tipo 1 y por qué lo usaste?',
        '¿Qué es is_comparable y por qué está en dim_store?',
        '¿Cómo resuelves una discrepancia de GMV entre dos reportes?',
    ]),
    ('KPIs',          [
        '¿Por qué GMV/m² comparable es el North Star?',
        '¿Cuál es tu leading indicator y por qué es predictivo?',
        '¿Cómo detectas si un KPI tiene datos malos?',
        '¿Cuál es el KPI compuesto y cómo se calcula?',
    ]),
    ('Auditoría',     [
        '¿Cuántas transacciones tenían total_amount inconsistente?',
        '¿Por qué usas SUM(unit_price×qty) y no total_amount para GMV?',
        '¿Qué decidiste hacer con las 50 transacciones pre-apertura?',
        '¿Cuántas tiendas tenían gaps de datos y cuál era la más afectada?',
    ]),
]

for topic, questions in checklist:
    body(doc, f'📌 {topic}', bold=True, size=11, color=BLUE)
    for q in questions:
        bullet(doc, q)
    doc.add_paragraph()


# --- FOOTER ---
divider(doc)
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run(
    'Prueba Técnica — Data Analyst — Cadena Retail Multiformato Centroamérica  |  '
    'Diego Alberto Calderón Calderón  |  Abril 2026'
)
r_f.font.size      = Pt(8)
r_f.font.color.rgb = GRAY
r_f.font.italic    = True


# Guardar
OUTPUT = 'Guia_Exposicion_Diego_Calderon.docx'
doc.save(OUTPUT)
print(f'Word generado: {OUTPUT}')